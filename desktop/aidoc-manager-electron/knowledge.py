"""本地知识索引：静默保存正文分段，并为 AI 秘书检索相关片段。

原文件始终留在用户资料库；索引只保存可重建的本地文字和出处。
"""

from __future__ import annotations

import hashlib
import re
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Any


MAX_DOCUMENT_CHARS = 5_000_000
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 180


def ensure_tables(conn: sqlite3.Connection) -> None:
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS knowledge_documents (
          document_id INTEGER PRIMARY KEY,
          source_sha256 TEXT NOT NULL DEFAULT '',
          status TEXT NOT NULL DEFAULT 'pending',
          extractor TEXT NOT NULL DEFAULT '',
          text_chars INTEGER NOT NULL DEFAULT 0,
          chunk_count INTEGER NOT NULL DEFAULT 0,
          content_hash TEXT NOT NULL DEFAULT '',
          error TEXT NOT NULL DEFAULT '',
          indexed_at TEXT NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS knowledge_chunks (
          id INTEGER PRIMARY KEY AUTOINCREMENT,
          document_id INTEGER NOT NULL,
          page_number INTEGER NOT NULL DEFAULT 0,
          chunk_index INTEGER NOT NULL DEFAULT 0,
          kind TEXT NOT NULL DEFAULT 'content',
          text TEXT NOT NULL,
          text_hash TEXT NOT NULL,
          created_at TEXT NOT NULL,
          UNIQUE(document_id, page_number, chunk_index, kind)
        )
        """
    )
    conn.execute("CREATE INDEX IF NOT EXISTS idx_knowledge_chunks_document ON knowledge_chunks(document_id, page_number)")
    has_fts = conn.execute(
        "SELECT 1 FROM sqlite_master WHERE type = 'table' AND name = 'knowledge_chunks_fts'"
    ).fetchone()
    if not has_fts:
        try:
            conn.execute(
                "CREATE VIRTUAL TABLE knowledge_chunks_fts USING fts5("
                "text, content='knowledge_chunks', content_rowid='id', tokenize='trigram')"
            )
        except sqlite3.OperationalError:
            conn.execute(
                "CREATE VIRTUAL TABLE knowledge_chunks_fts USING fts5("
                "text, content='knowledge_chunks', content_rowid='id', tokenize='unicode61')"
            )
    conn.executescript(
        """
        CREATE TRIGGER IF NOT EXISTS knowledge_chunks_ai AFTER INSERT ON knowledge_chunks BEGIN
          INSERT INTO knowledge_chunks_fts(rowid, text) VALUES (new.id, new.text);
        END;
        CREATE TRIGGER IF NOT EXISTS knowledge_chunks_ad AFTER DELETE ON knowledge_chunks BEGIN
          INSERT INTO knowledge_chunks_fts(knowledge_chunks_fts, rowid, text)
          VALUES ('delete', old.id, old.text);
        END;
        CREATE TRIGGER IF NOT EXISTS knowledge_chunks_au AFTER UPDATE ON knowledge_chunks BEGIN
          INSERT INTO knowledge_chunks_fts(knowledge_chunks_fts, rowid, text)
          VALUES ('delete', old.id, old.text);
          INSERT INTO knowledge_chunks_fts(rowid, text) VALUES (new.id, new.text);
        END;
        """
    )
    if not has_fts:
        conn.execute("INSERT INTO knowledge_chunks_fts(knowledge_chunks_fts) VALUES ('rebuild')")
    conn.commit()


def cleanup_removed_documents(conn: sqlite3.Connection) -> None:
    conn.execute("DELETE FROM knowledge_chunks WHERE document_id NOT IN (SELECT id FROM documents)")
    conn.execute("DELETE FROM knowledge_documents WHERE document_id NOT IN (SELECT id FROM documents)")
    conn.commit()


def pending_document_ids(conn: sqlite3.Connection, limit: int = 5000) -> list[int]:
    rows = conn.execute(
        """
        SELECT d.id
        FROM documents d
        LEFT JOIN knowledge_documents k ON k.document_id = d.id
        WHERE d.recognized_at IS NOT NULL
          AND (k.document_id IS NULL OR k.source_sha256 <> d.sha256)
        ORDER BY d.id ASC
        LIMIT ?
        """,
        (max(1, min(10000, int(limit))),),
    ).fetchall()
    return [int(row[0]) for row in rows]


def _normalize_text(value: str) -> str:
    text = str(value or "").replace("\x00", "").replace("\u2f8f", "行")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _read_text(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def extract_pages(path: Path) -> tuple[list[tuple[int, str]], str]:
    """返回逻辑页文字。扫描件没有文字层时返回空页列表。"""
    suffix = path.suffix.lower()
    pages: list[tuple[int, str]] = []
    if suffix == ".pdf":
        import fitz

        # 个别历史 PDF 的表单外观流或内容流损坏，MuPDF 会向控制台打印警告；
        # 索引仍可继续提取其它页面，因此静默这些底层诊断并由状态字段记录结果。
        fitz.TOOLS.mupdf_display_errors(False)
        document = fitz.open(str(path))
        try:
            for page_index, page in enumerate(document):
                text = _normalize_text(page.get_text())
                if text:
                    pages.append((page_index + 1, text))
        finally:
            document.close()
        return pages, "pdf_text_layer"
    if suffix == ".docx":
        from docx import Document

        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                value = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if value:
                    parts.append(value)
        text = _normalize_text("\n".join(parts))
        return ([(1, text)] if text else []), "docx"
    if suffix in {".xlsx", ".xlsm"}:
        import openpyxl

        workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        try:
            for sheet_index, worksheet in enumerate(workbook.worksheets):
                parts = [f"工作表：{worksheet.title}"]
                for row in worksheet.iter_rows(values_only=True):
                    value = " | ".join(str(cell) for cell in row if cell is not None)
                    if value:
                        parts.append(value)
                text = _normalize_text("\n".join(parts))
                if text:
                    pages.append((sheet_index + 1, text))
        finally:
            workbook.close()
        return pages, "spreadsheet"
    if suffix in {".txt", ".csv", ".md", ".markdown", ".json", ".xml", ".rtf", ".log"}:
        text = _normalize_text(_read_text(path))
        return ([(1, text)] if text else []), "text"
    return [], "no_local_text"


def _chunks(text: str) -> list[str]:
    text = _normalize_text(text)
    if not text:
        return []
    if len(text) <= CHUNK_SIZE:
        return [text]
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + CHUNK_SIZE)
        if end < len(text):
            boundary = max(text.rfind("\n", start + CHUNK_SIZE // 2, end), text.rfind("。", start + CHUNK_SIZE // 2, end))
            if boundary > start:
                end = boundary + 1
        value = text[start:end].strip()
        if value:
            chunks.append(value)
        if end >= len(text):
            break
        start = max(start + 1, end - CHUNK_OVERLAP)
    return chunks


def _metadata_text(row: sqlite3.Row) -> str:
    labels = [
        ("文件名", row["file_name"]),
        ("资料类型", row["document_type_label"]),
        ("公司", row["company_name"]),
        ("品牌", row["brand"]),
        ("证书或文号", row["certificate_no"]),
        ("出具方", row["issuer"]),
        ("签发日期", row["issued_at"]),
        ("有效期", row["expires_at"]),
        ("适用范围", row["applicable_scope"]),
        ("摘要", row["ai_summary"]),
        ("标签", row["tags"]),
    ]
    return "\n".join(f"{label}：{value}" for label, value in labels if str(value or "").strip())


def index_document(db_file: Path, document_id: int, *, force: bool = False) -> dict[str, Any]:
    conn = sqlite3.connect(str(db_file), timeout=30)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA busy_timeout = 30000")
    try:
        ensure_tables(conn)
        row = conn.execute("SELECT * FROM documents WHERE id = ?", (int(document_id),)).fetchone()
        if not row or not row["recognized_at"]:
            return {"indexed": False, "reason": "not_recognized"}
        existing = conn.execute(
            "SELECT source_sha256 FROM knowledge_documents WHERE document_id = ?", (int(document_id),)
        ).fetchone()
        if existing and str(existing[0] or "") == str(row["sha256"] or "") and not force:
            return {"indexed": False, "reason": "up_to_date"}
        managed = str(row["managed_path"] or "").strip()
        original = str(row["original_path"] or "").strip()
        path = Path(managed) if managed and Path(managed).exists() else (Path(original) if original else None)
        pages: list[tuple[int, str]] = []
        extractor = "source_missing"
        error = ""
        if path is not None and path.is_file():
            try:
                pages, extractor = extract_pages(path)
            except Exception as exc:
                extractor = "extract_failed"
                error = f"{type(exc).__name__}: {str(exc)[:300]}"
        else:
            error = "源文件不存在"

        metadata = _metadata_text(row)
        records: list[tuple[int, int, str, str]] = []
        if metadata:
            records.append((0, 0, "metadata", metadata))
        total_chars = 0
        truncated = False
        for page_number, page_text in pages:
            if total_chars >= MAX_DOCUMENT_CHARS:
                truncated = True
                break
            remaining = MAX_DOCUMENT_CHARS - total_chars
            page_text = page_text[:remaining]
            total_chars += len(page_text)
            for chunk_index, chunk in enumerate(_chunks(page_text)):
                records.append((page_number, chunk_index, "content", chunk))
        if truncated:
            error = (error + "；" if error else "") + "正文超过本地索引上限，已截断"
        status = "full" if any(record[2] == "content" for record in records) else "metadata_only"
        content_digest = hashlib.sha256(
            "\n".join(record[3] for record in records).encode("utf-8", "ignore")
        ).hexdigest()
        now = datetime.now().isoformat(timespec="seconds")
        conn.execute("BEGIN IMMEDIATE")
        conn.execute("DELETE FROM knowledge_chunks WHERE document_id = ?", (int(document_id),))
        for page_number, chunk_index, kind, text in records:
            conn.execute(
                """
                INSERT INTO knowledge_chunks
                  (document_id, page_number, chunk_index, kind, text, text_hash, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    int(document_id), page_number, chunk_index, kind, text,
                    hashlib.sha256(text.encode("utf-8", "ignore")).hexdigest(), now,
                ),
            )
        conn.execute(
            """
            INSERT INTO knowledge_documents
              (document_id, source_sha256, status, extractor, text_chars, chunk_count,
               content_hash, error, indexed_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(document_id) DO UPDATE SET
              source_sha256 = excluded.source_sha256,
              status = excluded.status,
              extractor = excluded.extractor,
              text_chars = excluded.text_chars,
              chunk_count = excluded.chunk_count,
              content_hash = excluded.content_hash,
              error = excluded.error,
              indexed_at = excluded.indexed_at
            """,
            (
                int(document_id), str(row["sha256"] or ""), status, extractor, total_chars,
                len(records), content_digest, error, now,
            ),
        )
        conn.commit()
        return {
            "indexed": True,
            "document_id": int(document_id),
            "status": status,
            "text_chars": total_chars,
            "chunk_count": len(records),
        }
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def _query_grams(query: str) -> list[str]:
    compact = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "", str(query or ""))
    if len(compact) < 3:
        return [compact] if compact else []
    synonym_groups = {
        "金额": ["合同金额", "总金额", "合同价款", "货款合计", "合计金额"],
        "付款": ["付款条件", "付款时间", "支付时间", "付款期限", "结算方式"],
        "开户": ["开户银行", "开户行", "银行名称"],
        "账号": ["银行账号", "银行账户", "收款账号"],
        "有效期": ["有效期限", "截止日期", "到期日期"],
        "法人": ["法定代表人", "负责人姓名"],
        "地址": ["注册地址", "经营地址", "交付地点"],
        "规格": ["规格型号", "产品型号"],
    }
    phrases = [compact]
    for trigger, values in synonym_groups.items():
        if trigger in compact:
            phrases.extend(values)
    grams = []
    seen: set[str] = set()
    for phrase in phrases:
        if 3 <= len(phrase) <= 10 and phrase not in seen:
            seen.add(phrase)
            grams.append(phrase)
        for size in (4, 3):
            for index in range(0, len(phrase) - size + 1):
                value = phrase[index:index + size]
                if value not in seen:
                    seen.add(value)
                    grams.append(value)
                if len(grams) >= 60:
                    return grams
    return grams


def search(conn: sqlite3.Connection, query: str, limit: int = 12) -> list[dict[str, Any]]:
    ensure_tables(conn)
    grams = _query_grams(query)
    if not grams:
        return []
    match_query = " OR ".join('"' + gram.replace('"', '""') + '"' for gram in grams)
    try:
        rows = conn.execute(
            """
            SELECT c.id, c.document_id, c.page_number, c.kind, c.text,
                   d.file_name, d.document_type_label, d.company_name,
                   bm25(knowledge_chunks_fts) AS lexical_rank
            FROM knowledge_chunks_fts
            JOIN knowledge_chunks c ON c.id = knowledge_chunks_fts.rowid
            JOIN documents d ON d.id = c.document_id
            WHERE knowledge_chunks_fts MATCH ?
            ORDER BY lexical_rank ASC
            LIMIT 80
            """,
            (match_query,),
        ).fetchall()
    except sqlite3.OperationalError:
        token = max(grams, key=len)
        rows = conn.execute(
            """
            SELECT c.id, c.document_id, c.page_number, c.kind, c.text,
                   d.file_name, d.document_type_label, d.company_name, 0 AS lexical_rank
            FROM knowledge_chunks c
            JOIN documents d ON d.id = c.document_id
            WHERE c.text LIKE ?
            LIMIT 80
            """,
            (f"%{token}%",),
        ).fetchall()
    query_set = set(grams)
    ranked: list[tuple[float, sqlite3.Row]] = []
    for row in rows:
        compact_text = re.sub(r"\s+", "", str(row["text"] or ""))
        overlap = sum(1 for gram in query_set if gram in compact_text)
        score = overlap / max(1, len(query_set))
        if str(row["kind"] or "") == "metadata":
            score += 0.08
        try:
            score += max(0.0, min(0.2, -float(row["lexical_rank"] or 0) / 50.0))
        except (TypeError, ValueError):
            pass
        ranked.append((score, row))
    ranked.sort(key=lambda item: item[0], reverse=True)
    results: list[dict[str, Any]] = []
    per_document: dict[int, int] = {}
    for score, row in ranked:
        document_id = int(row["document_id"])
        if per_document.get(document_id, 0) >= 3:
            continue
        per_document[document_id] = per_document.get(document_id, 0) + 1
        results.append({
            "document_id": document_id,
            "file_name": str(row["file_name"] or ""),
            "type_label": str(row["document_type_label"] or ""),
            "company": str(row["company_name"] or ""),
            "page_number": int(row["page_number"] or 0),
            "kind": str(row["kind"] or "content"),
            "text": str(row["text"] or "")[:1800],
            "score": round(score, 4),
        })
        if len(results) >= max(1, min(30, int(limit))):
            break
    return results
